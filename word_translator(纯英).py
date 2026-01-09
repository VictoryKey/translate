#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档中文翻译成英文脚本
使用Google翻译API批量翻译Word文档内容
"""

import os
import time
import logging
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from deep_translator import GoogleTranslator

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class WordTranslator:
    """Word文档翻译器类"""
    
    def __init__(self, source_lang='zh-CN', target_lang='en'):
        """
        初始化翻译器
        
        Args:
            source_lang: 源语言 (默认中文)
            target_lang: 目标语言 (默认英文)
        """
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.translator = GoogleTranslator(source=source_lang, target=target_lang)
        
    def translate_text(self, text, retry_count=3):
        """
        翻译文本，支持重试机制
        
        Args:
            text: 要翻译的文本
            retry_count: 重试次数
            
        Returns:
            翻译后的文本
        """
        if not text or not text.strip():
            return text
            
        # 清理文本，移除多余的空白字符
        text = ' '.join(text.split())
        
        for attempt in range(retry_count):
            try:
                # Google翻译每次限制5000字符
                if len(text) > 4500:
                    # 分段翻译
                    chunks = [text[i:i+4500] for i in range(0, len(text), 4500)]
                    translated_chunks = []
                    for chunk in chunks:
                        result = self.translator.translate(chunk)
                        if result:
                            translated_chunks.append(result)
                        time.sleep(0.5)  # 避免请求过快
                    return ' '.join(translated_chunks)
                else:
                    result = self.translator.translate(text)
                    return result if result else text
                    
            except Exception as e:
                logger.warning(f"翻译失败 (尝试 {attempt + 1}/{retry_count}): {e}")
                if attempt < retry_count - 1:
                    time.sleep(2 * (attempt + 1))  # 递增等待时间
                else:
                    logger.error(f"翻译最终失败: {text[:50]}...")
                    return text
                    
    def translate_document(self, input_path, output_path=None, batch_size=50):
        """
        翻译整个Word文档
        
        Args:
            input_path: 输入文件路径
            output_path: 输出文件路径 (默认在原文件同级目录生成新文件)
            batch_size: 每次处理的段落数量
            
        Returns:
            输出文件路径
        """
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"文件不存在: {input_path}")
            
        if output_path is None:
            # 生成输出文件名
            base, ext = os.path.splitext(input_path)
            output_path = f"{base}_translated{ext}"
            
        logger.info(f"开始翻译文档: {input_path}")
        logger.info(f"输出文件: {output_path}")
        
        # 读取源文档
        doc = Document(input_path)
        
        # 获取所有段落
        paragraphs = doc.paragraphs
        total_paragraphs = len(paragraphs)
        
        logger.info(f"文档共有 {total_paragraphs} 个段落")
        
        # 翻译进度统计
        translated_count = 0
        skip_count = 0
        
        # 批量处理段落
        for i, paragraph in enumerate(paragraphs):
            # 显示进度
            if (i + 1) % 10 == 0 or i == 0:
                logger.info(f"进度: {i + 1}/{total_paragraphs} 段落")
            
            # 跳过空段落和只有空白的段落
            if not paragraph.text or not paragraph.text.strip():
                skip_count += 1
                continue
                
            # 检查是否包含中文字符
            has_chinese = any(ord(char) > 127 for char in paragraph.text)
            if not has_chinese:
                # 没有中文，直接跳过（也可以选择保留或翻译）
                continue
                
            # 翻译段落
            original_text = paragraph.text
            translated_text = self.translate_text(original_text)
            
            if translated_text and translated_text != original_text:
                # 保存原始样式
                style = paragraph.style
                
                # 清空段落并添加翻译后的文本
                paragraph.clear()
                run = paragraph.add_run(translated_text)
                
                # 应用原始样式
                if style:
                    paragraph.style = style
                    # 应用字体样式
                    if style.font.size:
                        run.font.size = style.font.size
                    if style.font.name:
                        run.font.name = style.font.name
                        
                translated_count += 1
                
            # 添加延迟，避免请求过快
            time.sleep(0.3)
            
        # 处理表格
        tables = doc.tables
        total_tables = len(tables)
        
        if total_tables > 0:
            logger.info(f"开始处理 {total_tables} 个表格")
            
            for table_idx, table in enumerate(tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        if cell.text and cell.text.strip():
                            has_chinese = any(ord(char) > 127 for char in cell.text)
                            if has_chinese:
                                original_text = cell.text
                                translated_text = self.translate_text(original_text)
                                
                                if translated_text and translated_text != original_text:
                                    cell.text = translated_text
                                    translated_count += 1
                                    
                                time.sleep(0.2)
                
                if (table_idx + 1) % 5 == 0:
                    logger.info(f"表格进度: {table_idx + 1}/{total_tables}")
        
        # 保存文档
        doc.save(output_path)
        
        logger.info(f"翻译完成!")
        logger.info(f"翻译段落/表格单元格数: {translated_count}")
        logger.info(f"跳过数: {skip_count}")
        logger.info(f"输出文件: {output_path}")
        
        return output_path

def main():
    """主函数"""
    # 示例用法
    input_file = "C:\\Users\\Bamb00\\Desktop\\Penetration Test Report For BMW IDCEvo V2.0（中文）.docx"  # 输入文件路径
    
    # 创建翻译器实例
    translator = WordTranslator(source_lang='zh-CN', target_lang='en')
    
    try:
        # 翻译文档
        output_file = translator.translate_document(input_file)
        print(f"\n翻译完成! 输出文件: {output_file}")
        
    except FileNotFoundError as e:
        print(f"错误: {e}")
    except Exception as e:
        print(f"翻译过程中发生错误: {e}")

if __name__ == "__main__":
    main()
