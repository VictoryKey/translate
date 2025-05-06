"""
Word文档智能翻译工具
功能：自动翻译中文Word文档为英文，保留原有格式，支持缓存、重试、多线程等特性
"""

# 标准库导入
import hashlib    # 用于生成MD5签名
import random     # 生成随机数
import requests   # 发送HTTP请求
import os         # 文件路径操作
import json       # 缓存文件读写
import threading  # 线程锁
import argparse   # 命令行参数解析
import time       # 时间控制
import re         # 正则表达式处理
from concurrent.futures import ThreadPoolExecutor, as_completed  # 线程池管理
from docx import Document  # Word文档操作
from docx.shared import RGBColor  # 颜色设置
from docx.oxml import OxmlElement  # 操作Word XML元素
from docx.text.paragraph import Paragraph  # 段落操作
from docx.oxml.ns import qn  # XML命名空间处理
from tqdm import tqdm  # 进度条显示

# ========== 全局配置 ==========
appid = '20241223002235907'      # 百度翻译API应用ID（需自行申请）
secretKey = 'FaoqgT8hOE5JLFoySt9S'  # 百度翻译API密钥
url = 'https://fanyi-api.baidu.com/api/trans/vip/translate'  # API端点
MAX_RETRIES = 10    # 最大重试次数
MAX_WORKERS = 4     # 最大线程数
CACHE_FILE = 'translation_cache.json'  # 缓存文件路径

# ========== 初始化缓存与锁 ==========
# 加载翻译缓存（如果存在）
if os.path.exists(CACHE_FILE):
    with open(CACHE_FILE, 'r', encoding='utf-8') as f:
        translation_cache = json.load(f)
else:
    translation_cache = {}

# 创建线程锁
cache_lock = threading.Lock()   # 缓存访问锁
print_lock = threading.Lock()   # 控制台输出锁

def highlight_paragraph(paragraph):
    """为段落添加黄色高亮效果（用于标记翻译失败的段落）"""
    for run in paragraph.runs:
        # 创建XML高亮元素
        rPr = run._element.get_or_add_rPr()
        highlight = OxmlElement('w:highlight')
        highlight.set(qn('w:val'), 'yellow')  # 设置高亮颜色为黄色
        rPr.append(highlight)

def save_cache():
    """将缓存保存到文件（线程安全）"""
    with cache_lock:
        with open(CACHE_FILE, 'w', encoding='utf-8') as f:
            json.dump(translation_cache, f, ensure_ascii=False, indent=2)

def baidu_translate(text, attempt):
    """
    调用百度翻译API进行单次翻译尝试
    :param text: 待翻译文本
    :param attempt: 当前尝试次数（用于错误提示）
    :return: 翻译结果或None
    """
    # 生成随机盐值
    salt = random.randint(32768, 65536)
    # 构造签名
    sign_str = appid + text + str(salt) + secretKey
    sign = hashlib.md5(sign_str.encode()).hexdigest()
    
    # 请求参数
    params = {
        'q': text,
        'from': 'zh',
        'to': 'en',
        'appid': appid,
        'salt': salt,
        'sign': sign
    }
    
    try:
        # 发送请求
        r = requests.get(url, params=params, timeout=10)
        r.raise_for_status()  # 检查HTTP状态码
        
        result = r.json()
        # 验证响应结构
        if 'trans_result' not in result:
            raise ValueError(f"Invalid response: {result}")
            
        return result['trans_result'][0]['dst']
    except Exception as e:
        # 输出错误信息（线程安全）
        with print_lock:
            tqdm.write(f"⚠️ 尝试 {attempt} 失败: {text[:20]}... ({str(e)})")
        return None

def baidu_translate_cached(text):
    """
    带缓存和重试机制的翻译函数
    :param text: 待翻译文本
    :return: 翻译结果或None
    """
    # 检查缓存（如果未禁用缓存）
    if not args.nocache:
        with cache_lock:
            if text in translation_cache:
                return translation_cache[text]

    # 特殊处理：直接映射"回归测试"
    if '回归测试' in text:
        translation_cache[text] = 'Regression testing'
        return 'Regression testing'

    # 处理带编号的文本（如"1. 测试内容"）
    # 使用正则表达式匹配编号结构
    match = re.match(r'^(\d+(?:\.\d+)*)([\s:：，,]+)(.+)$', text)
    if match:
        prefix = match.group(1)    # 编号部分（如"1.2.3"）
        separator = match.group(2) # 分隔符（如":"或空格）
        rest = match.group(3)      # 实际需要翻译的内容

        # 仅翻译内容部分
        translated_rest = None
        for attempt in range(1, MAX_RETRIES + 1):
            translated_rest = baidu_translate(rest.strip(), attempt)
            if translated_rest:
                break
            # 指数退避策略：2^attempt秒 + 随机延迟（避免请求风暴）
            time.sleep(2 ** min(attempt, 5) + random.random())

        if translated_rest:
            # 组合编号和翻译结果
            result = f"{prefix}{separator}{translated_rest}"
            with cache_lock:
                translation_cache[text] = result
            return result
        else:
            return None

    # 常规翻译流程
    translated = None
    for attempt in range(1, MAX_RETRIES + 1):
        translated = baidu_translate(text, attempt)
        if translated:
            # 重试成功提示
            if attempt > 1:
                with print_lock:
                    tqdm.write(f"✅ 重试后成功: {text[:30]} (第 {attempt} 次)")
            # 更新缓存
            with cache_lock:
                translation_cache[text] = translated
            return translated
        # 指数退避策略
        time.sleep(2 ** min(attempt, 5) + random.random())

    return None

def split_into_sentences(text):
    """
    将文本分割成句子
    :param text: 原始文本
    :return: 句子列表
    """
    # 使用正则表达式分割句子：
    # - 在句号、感叹号、问号、分号后分割
    # - 在换行符前分割
    # - 在数字编号（如"1."）前分割
    return [s.strip() for s in re.split(r'(?<=[。！？；])|(?=\n)|(?=\d+\.)', text) if s.strip()]

def is_title(paragraph):
    """
    判断段落是否为标题（根据字体大小）
    :param paragraph: Word段落对象
    :return: 是否是标题
    """
    if paragraph.runs:
        font_size = paragraph.runs[0].font.size
        # 判断字体大小是否接近16pt（标题常见大小）
        if font_size and abs(font_size.pt - 16) < 1:
            return True
    return False
    # if paragraph.runs:
    #     font_size = paragraph.runs[0].font.size
    #     # 判断字体大小是否接近16pt（标题常见大小）
    #     if font_size == '三号':
    #         return True
    # return False

def process_paragraph(paragraph, pbar, stats):
    """
    处理单个段落（线程执行单元）
    :param paragraph: Word段落对象
    :param pbar: 进度条对象
    :param stats: 统计信息字典
    """
    text = paragraph.text.strip()
    if not text:
        pbar.update(1)
        return

    # 分割文本为多个句子
    parts = split_into_sentences(text)
    combined_result = []  # 组合后的翻译结果
    success = True        # 是否全部翻译成功

    # 逐句翻译
    for part in parts:
        translated = baidu_translate_cached(part)
        if translated:
            combined_result.append(translated)
        else:
            combined_result.append(f"[翻译失败: {part}]")
            success = False

    # 处理标题的特殊情况
    if is_title(paragraph):
        # 在原文后追加翻译（保留标题格式）
        paragraph.add_run(' ' + ' '.join(combined_result))
    else:
        # 创建新段落插入翻译结果
        new_para = insert_paragraph_after(paragraph)
        run = new_para.add_run(' '.join(combined_result))
        
        # 复制原段落样式
        if paragraph.runs:
            copy_run_style(paragraph.runs[0], run)
        
        # 标记部分失败的段落
        if not success:
            run.font.color.rgb = RGBColor(255, 255, 0)  # 设置文字为黄色
            highlight_paragraph(new_para)  # 添加背景高亮

    # 更新统计信息
    if success:
        stats['success'] += 1
    else:
        stats['failed'] += 1

    save_cache()
    pbar.update(1)

def insert_paragraph_after(paragraph):
    """在当前段落之后插入新段落"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)

def copy_run_style(source_run, target_run):
    """复制文字样式到目标Run"""
    try:
        # 复制基础样式
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        target_run.font.name = source_run.font.name
        target_run.font.size = source_run.font.size
        
        # 处理中文字体设置
        if target_run._element.rPr is None:
            target_run._element.get_or_add_rPr()
        if source_run.font.name:
            target_run._element.rPr.rFonts.set(qn('w:eastAsia'), source_run.font.name)
    except Exception as e:
        pass

def translate_paragraphs(paragraphs):
    """执行多线程翻译"""
    stats = {'success': 0, 'failed': 0}
    # 计算实际需要翻译的段落数（过滤空段落）
    total = len([p for p in paragraphs if p.text.strip()])

    with tqdm(total=total, desc="翻译进度", ncols=80) as pbar:
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            futures = []
            # 提交翻译任务
            for p in paragraphs:
                if p.text.strip():
                    future = executor.submit(process_paragraph, p, pbar, stats)
                    futures.append(future)

            # 等待任务完成
            for future in as_completed(futures):
                try:
                    future.result()
                except Exception as e:
                    with print_lock:
                        tqdm.write(f"⚠️ 线程错误: {str(e)}")

    return stats

def run_translate(file_path):
    """主执行函数"""
    global translation_cache
    
    # 处理缓存相关参数
    if args.resetcache and os.path.exists(CACHE_FILE):
        os.remove(CACHE_FILE)
        translation_cache = {}
        print("🗑 已清空缓存文件")
    elif not os.path.exists(CACHE_FILE):
        translation_cache = {}
    elif not args.nocache:
        with open(CACHE_FILE, 'r', encoding='utf-8') as f:
            translation_cache = json.load(f)

    # 加载Word文档
    print(f"📄 正在处理文档：{file_path}")
    doc = Document(file_path)

    # 收集所有需要翻译的段落
    all_paragraphs = []
    all_paragraphs.extend(doc.paragraphs)  # 添加正文段落
    
    # 添加表格中的段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    print(f"📝 发现可翻译段落: {len(all_paragraphs)} 个")
    
    # 执行翻译
    stats = translate_paragraphs(all_paragraphs)

    # 保存结果文档
    base_name, ext = os.path.splitext(file_path)
    new_name = base_name + "_translate" + ext
    doc.save(new_name)
    print(f"\n✅ 翻译完成 - 成功: {stats['success']}, 失败: {stats['failed']}")
    print(f"📁 文件已保存为：{new_name}")

if __name__ == "__main__":
    # 命令行参数解析
    parser = argparse.ArgumentParser(description="Word文档智能翻译工具")
    parser.add_argument('-w', '--wordfile', required=True, help='Word文件路径')
    parser.add_argument('--nocache', action='store_true', help='忽略缓存并强制重新翻译')
    parser.add_argument('--resetcache', action='store_true', help='清空缓存后再翻译')
    args = parser.parse_args()
    
    # 启动翻译流程
    run_translate(args.wordfile)